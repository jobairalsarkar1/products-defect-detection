"""
evaluate_and_visualize.py

Usage:
    python evaluate_and_visualize.py

What it does:
- Finds all *_best_model.keras files under ./saved_models/*
- Loads each model
- Finds images under ./test_images/good and ./test_images/bad (keeps alphabetical order)
- Predicts each image, supports both single-output sigmoid and 2-class softmax outputs
- Creates a large composite PNG: one row per model, 8 columns (4 good then 4 bad)
- Saves results image to ./results/model_comparison.png
- Saves a Word report to ./results/report.docx
"""

import os
import sys
from pathlib import Path
from typing import List, Tuple, Dict

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from tensorflow.keras.models import load_model
from tensorflow.keras.preprocessing import image as keras_image
from docx import Document
from docx.shared import Inches
from tqdm import tqdm

# =========== Config ==============
ROOT = Path(__file__).resolve().parent
SAVED_MODELS_DIR = ROOT / "saved_models"
TEST_IMAGES_DIR = ROOT / "test_images"   # contains 'good' and 'bad' subfolders
RESULTS_DIR = ROOT / "results"
IMG_SIZE = (224, 224)  # change if your models expect different
CLASS_MAPPING = {0: "bad", 1: "good"}  # assumed mapping
CONF_THRESHOLD = 0.5  # threshold for binary sigmoid outputs
TILE_W = 320
TILE_H = 260
# None -> use default PIL font. Set to a .ttf path for nicer text.
FONT_PATH = None
# ==================================

RESULTS_DIR.mkdir(parents=True, exist_ok=True)


def find_best_models(saved_models_dir: Path) -> List[Path]:
    """Search saved_models/* for files with *_best_model.keras (case-insensitive)."""
    found = []
    for sub in saved_models_dir.iterdir():
        if sub.is_dir():
            for f in sub.iterdir():
                if f.is_file() and f.name.lower().endswith("_best_model.keras"):
                    found.append(f)
    # sort for deterministic order
    found.sort()
    return found


def load_and_prepare_image(img_path: Path, target_size: Tuple[int, int]) -> np.ndarray:
    """Load image and return tensor ready for model.predict: shape (1,H,W,3)."""
    img = keras_image.load_img(
        str(img_path), target_size=target_size, color_mode="rgb")
    arr = keras_image.img_to_array(img).astype("float32") / 255.0
    return np.expand_dims(arr, axis=0)


def softmax(x):
    e = np.exp(x - np.max(x))
    return e / e.sum(axis=-1, keepdims=True)


def predict_with_model(model, img_tensor: np.ndarray) -> Tuple[int, float, np.ndarray]:
    """
    Return (predicted_class_index, confidence_for_predicted_class, raw_output_array)
    Handles:
      - single-sigmoid output => one float: interprets >0.5 as class 1
      - 2-element softmax => argmax
      - >2 classes => argmax
    """
    out = model.predict(img_tensor, verbose=0)
    out = np.asarray(out)
    # flatten if necessary: e.g., (1,1) or (1,)
    if out.size == 1:
        prob = float(out.flatten()[0])
        if prob > CONF_THRESHOLD:
            return 1, prob, out
        else:
            return 0, 1.0 - prob, out
    # if shape (1,2) or (2,) treat as softmax-like
    if out.ndim == 2 and out.shape[1] >= 2:
        probs = out[0]
        # if not normalized, softmax
        if not np.isclose(probs.sum(), 1.0, atol=1e-3):
            probs = softmax(probs)
        idx = int(np.argmax(probs))
        conf = float(probs[idx])
        return idx, conf, probs
    # fallback: flatten and argmax
    flat = out.flatten()
    idx = int(np.argmax(flat))
    # attempt to get confidence if interpretable
    if flat.sum() > 0:
        conf = float(flat[idx] / flat.sum())
    else:
        conf = float(flat[idx])
    return idx, conf, flat


def build_visual_grid(models_info, image_paths_good, image_paths_bad, out_path: Path):
    """
    models_info: list of tuples (model_name, predictions_list)
      predictions_list: list of dicts with keys:
         'img_path', 'true_label', 'pred_label', 'confidence', 'correct'
    """
    rows = len(models_info)
    cols = len(image_paths_good) + len(image_paths_bad)
    W = TILE_W * cols
    H = TILE_H * rows
    canvas = Image.new("RGB", (W, H), color=(245, 245, 245))

    # font handling
    if FONT_PATH and Path(FONT_PATH).exists():
        font_title = ImageFont.truetype(FONT_PATH, 18)
        font_small = ImageFont.truetype(FONT_PATH, 14)
        font_big = ImageFont.truetype(FONT_PATH, 20)
    else:
        font_title = ImageFont.load_default()
        font_small = ImageFont.load_default()
        font_big = ImageFont.load_default()

    draw = ImageDraw.Draw(canvas)

    for r, (model_name, preds) in enumerate(models_info):
        # draw model name left of row (we'll draw inside first tile)
        for c, pred in enumerate(preds):
            x0 = c * TILE_W
            y0 = r * TILE_H
            box = (x0, y0, x0 + TILE_W, y0 + TILE_H)
            # load thumbnail
            try:
                thumb = Image.open(pred['img_path']).convert("RGB")
                thumb = thumb.resize((TILE_W, int(TILE_H*0.6)))
            except Exception:
                thumb = Image.new(
                    "RGB", (TILE_W, int(TILE_H*0.6)), color=(200, 200, 200))
            canvas.paste(thumb, (x0, y0))
            # draw labels area below the image
            text_x = x0 + 6
            text_y = y0 + int(TILE_H*0.6) + 6
            # filename (short)
            fname = Path(pred['img_path']).name
            draw.text((text_x, text_y), fname,
                      font=font_small, fill=(10, 10, 10))
            # true / predicted / confidence
            text_y += 16
            draw.text(
                (text_x, text_y), f"True: {pred['true_label']}", font=font_small, fill=(10, 10, 10))
            text_y += 16
            # predicted and confidence
            pred_text = f"Pred: {pred['pred_label']}  ({pred['confidence']*100:.1f}%)"
            draw.text((text_x, text_y), pred_text,
                      font=font_small, fill=(10, 10, 10))
            # correct tick
            mark = "✅" if pred['correct'] else "❌"
            draw.text((x0 + TILE_W - 36, y0 + int(TILE_H*0.6) + 8),
                      mark, font=font_big, fill=(10, 10, 10))
            # draw thin separator
            draw.rectangle([x0, y0, x0+TILE_W, y0+TILE_H],
                           outline=(230, 230, 230))

        # write model name as a label on the leftmost margin of the row (over first tile)
        model_label_x = 6
        model_label_y = r * TILE_H + int(TILE_H*0.6) - 22
        draw.rectangle([0 + 6, model_label_y - 2, 6 + 260,
                       model_label_y + 20], fill=(255, 255, 255))
        draw.text((model_label_x + 8, model_label_y),
                  model_name, font=font_title, fill=(0, 0, 0))

    canvas.save(out_path)
    print(f"Saved composite image to: {out_path}")


def create_docx_report(models_info, out_path: Path, image_png_path: Path):
    doc = Document()
    doc.add_heading("Model Comparison Report", level=1)
    doc.add_paragraph(f"Generated at: {os.path.abspath(out_path)}")
    doc.add_paragraph(
        "Summary of model predictions on test_images (4 good + 4 bad).")

    # overall table (model, accuracy)
    doc.add_heading("Accuracy Summary", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model'
    hdr_cells[1].text = 'Correct / Total'
    hdr_cells[2].text = 'Accuracy (%)'
    for model_name, preds in models_info:
        correct = sum(1 for p in preds if p['correct'])
        total = len(preds)
        acc = (correct / total) * 100 if total > 0 else 0.0
        row_cells = table.add_row().cells
        row_cells[0].text = model_name
        row_cells[1].text = f"{correct} / {total}"
        row_cells[2].text = f"{acc:.2f}"

    doc.add_heading("Composite Visualization", level=2)
    doc.add_paragraph(
        "Below is the combined visualization (one row per model, 8 images each).")
    doc.add_picture(str(image_png_path), width=Inches(6.5))

    doc.add_heading("Detailed Per-Image Results", level=2)
    for model_name, preds in models_info:
        doc.add_heading(model_name, level=3)
        t = doc.add_table(rows=1, cols=5)
        h = t.rows[0].cells
        h[0].text = "Image"
        h[1].text = "True"
        h[2].text = "Predicted"
        h[3].text = "Confidence"
        h[4].text = "Correct"
        for p in preds:
            row = t.add_row().cells
            row[0].text = Path(p['img_path']).name
            row[1].text = p['true_label']
            row[2].text = p['pred_label']
            row[3].text = f"{p['confidence']*100:.1f}%"
            row[4].text = "Yes" if p['correct'] else "No"
        doc.add_paragraph("")

    doc.save(out_path)
    print(f"Saved report to: {out_path}")


def main():
    # 1) find models
    model_files = find_best_models(SAVED_MODELS_DIR)
    if not model_files:
        print("No *_best_model.keras files found under", SAVED_MODELS_DIR)
        sys.exit(1)

    print("Found models:")
    for m in model_files:
        print("  ", m)

    # 2) gather test images
    good_dir = TEST_IMAGES_DIR / "good"
    bad_dir = TEST_IMAGES_DIR / "bad"
    if not good_dir.exists() or not bad_dir.exists():
        print(
            "Make sure test_images/good and test_images/bad exist under:", TEST_IMAGES_DIR)
        sys.exit(1)
    # keep alphabetical order
    image_paths_good = sorted([p for p in good_dir.iterdir(
    ) if p.suffix.lower() in [".jpg", ".jpeg", ".png"]])
    image_paths_bad = sorted([p for p in bad_dir.iterdir(
    ) if p.suffix.lower() in [".jpg", ".jpeg", ".png"]])

    if len(image_paths_good) == 0 or len(image_paths_bad) == 0:
        print("No images found in good/ or bad/ folders.")
        sys.exit(1)

    # 3) for each model, load model and predict
    models_info = []
    for model_path in model_files:
        model_name = model_path.parent.name if model_path.parent else model_path.stem
        print(f"\nLoading model: {model_name} ({model_path})")
        try:
            model = load_model(str(model_path))
        except Exception as e:
            print(f"Failed to load {model_path}: {e}")
            continue

        preds_for_model = []
        # iterate good then bad to ensure columns order is 4 good then 4 bad
        for p in list(image_paths_good) + list(image_paths_bad):
            img_tensor = load_and_prepare_image(p, IMG_SIZE)
            pred_idx, conf, raw = predict_with_model(model, img_tensor)
            pred_label = CLASS_MAPPING.get(pred_idx, str(pred_idx))
            true_label = "good" if p.parent.name.lower() == "good" else "bad"
            correct = (pred_label.lower() == true_label.lower())
            preds_for_model.append({
                "img_path": str(p),
                "true_label": true_label,
                "pred_label": pred_label,
                "confidence": float(conf),
                "correct": bool(correct),
            })
            print(
                f"  {p.name} -> {pred_label} ({conf*100:.1f}%)  {'OK' if correct else 'WRONG'}")

        models_info.append((model_name, preds_for_model))

    # 4) build composite image
    out_png = RESULTS_DIR / "model_comparison.png"
    build_visual_grid(models_info, image_paths_good, image_paths_bad, out_png)

    # 5) create docx report
    out_docx = RESULTS_DIR / "report.docx"
    create_docx_report(models_info, out_docx, out_png)


if __name__ == "__main__":
    main()
