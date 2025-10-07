"""
evaluate_and_visualize.py (v2)
Generates a clean visualization: 2 good + 2 bad images per model.
"""

import os
import sys
from pathlib import Path
from typing import List, Tuple, Dict

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from tensorflow.keras.models import load_model
from tensorflow.keras.preprocessing import image as keras_image
from docx import Document
from docx.shared import Inches
from tqdm import tqdm

# =========== Config ==============
ROOT = Path(__file__).resolve().parent
SAVED_MODELS_DIR = ROOT / "saved_models"
TEST_IMAGES_DIR = ROOT / "test_images"   # contains 'good' and 'bad'
RESULTS_DIR = ROOT / "results"
IMG_SIZE = (224, 224)
CLASS_MAPPING = {0: "bad", 1: "good"}
CONF_THRESHOLD = 0.5
TILE_W = 320
TILE_H = 260
FONT_PATH = None  # or e.g. "C:/Windows/Fonts/arial.ttf"
NUM_GOOD = 2
NUM_BAD = 2
# ==================================

RESULTS_DIR.mkdir(parents=True, exist_ok=True)

def find_best_models(saved_models_dir: Path) -> List[Path]:
    found = []
    for sub in saved_models_dir.iterdir():
        if sub.is_dir():
            for f in sub.iterdir():
                if f.is_file() and f.name.lower().endswith("_best_model.keras"):
                    found.append(f)
    found.sort()
    return found

def load_and_prepare_image(img_path: Path, target_size: Tuple[int,int]) -> np.ndarray:
    img = keras_image.load_img(str(img_path), target_size=target_size, color_mode="rgb")
    arr = keras_image.img_to_array(img).astype("float32") / 255.0
    return np.expand_dims(arr, axis=0)

def softmax(x):
    e = np.exp(x - np.max(x))
    return e / e.sum(axis=-1, keepdims=True)

def predict_with_model(model, img_tensor: np.ndarray) -> Tuple[int, float, np.ndarray]:
    out = model.predict(img_tensor, verbose=0)
    out = np.asarray(out)
    if out.size == 1:
        prob = float(out.flatten()[0])
        return (1, prob, out) if prob > CONF_THRESHOLD else (0, 1.0 - prob, out)
    if out.ndim == 2 and out.shape[1] >= 2:
        probs = out[0]
        if not np.isclose(probs.sum(), 1.0, atol=1e-3):
            probs = softmax(probs)
        idx = int(np.argmax(probs))
        conf = float(probs[idx])
        return idx, conf, probs
    flat = out.flatten()
    idx = int(np.argmax(flat))
    conf = float(flat[idx] / flat.sum()) if flat.sum() > 0 else float(flat[idx])
    return idx, conf, flat

def build_visual_grid(models_info, out_path: Path):
    rows = len(models_info)
    cols = NUM_GOOD + NUM_BAD
    W = TILE_W * cols
    H = TILE_H * rows
    canvas = Image.new("RGB", (W, H), color=(245, 245, 245))

    if FONT_PATH and Path(FONT_PATH).exists():
        font_title = ImageFont.truetype(FONT_PATH, 18)
        font_small = ImageFont.truetype(FONT_PATH, 14)
        font_big = ImageFont.truetype(FONT_PATH, 20)
    else:
        font_title = ImageFont.load_default()
        font_small = ImageFont.load_default()
        font_big = ImageFont.load_default()

    draw = ImageDraw.Draw(canvas)

    for r, (model_name, preds) in enumerate(models_info):
        for c, pred in enumerate(preds):
            x0 = c * TILE_W
            y0 = r * TILE_H
            try:
                thumb = Image.open(pred['img_path']).convert("RGB")
                thumb = thumb.resize((TILE_W, int(TILE_H*0.6)))
            except Exception:
                thumb = Image.new("RGB", (TILE_W, int(TILE_H*0.6)), color=(200,200,200))
            canvas.paste(thumb, (x0, y0))
            text_x = x0 + 6
            text_y = y0 + int(TILE_H*0.6) + 6
            fname = Path(pred['img_path']).name
            draw.text((text_x, text_y), fname, font=font_small, fill=(10,10,10))
            text_y += 16
            draw.text((text_x, text_y), f"True: {pred['true_label']}", font=font_small, fill=(10,10,10))
            text_y += 16
            pred_text = f"Pred: {pred['pred_label']} ({pred['confidence']*100:.1f}%)"
            draw.text((text_x, text_y), pred_text, font=font_small, fill=(10,10,10))
            mark = "✅" if pred['correct'] else "❌"
            draw.text((x0 + TILE_W - 36, y0 + int(TILE_H*0.6) + 8), mark, font=font_big, fill=(10,10,10))
            draw.rectangle([x0, y0, x0+TILE_W, y0+TILE_H], outline=(230,230,230))

        model_label_x = 8
        model_label_y = r * TILE_H + int(TILE_H*0.6) - 20
        draw.rectangle([model_label_x - 2, model_label_y - 4, model_label_x + 180, model_label_y + 24], fill=(255,255,255))
        draw.text((model_label_x, model_label_y), model_name, font=font_title, fill=(0,0,0))

    canvas.save(out_path)
    print(f"✅ Saved composite image to: {out_path}")

def create_docx_report(models_info, out_path: Path, image_png_path: Path):
    doc = Document()
    doc.add_heading("Model Comparison Report", level=1)
    doc.add_paragraph("Each model is tested on 4 images (2 good + 2 bad).")

    doc.add_heading("Accuracy Summary", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = 'Model'
    hdr[1].text = 'Correct / Total'
    hdr[2].text = 'Accuracy (%)'

    for model_name, preds in models_info:
        correct = sum(p['correct'] for p in preds)
        total = len(preds)
        acc = (correct / total) * 100
        row = table.add_row().cells
        row[0].text = model_name
        row[1].text = f"{correct} / {total}"
        row[2].text = f"{acc:.2f}"

    doc.add_heading("Composite Visualization", level=2)
    doc.add_picture(str(image_png_path), width=Inches(6.5))

    doc.save(out_path)
    print(f"✅ Saved report to: {out_path}")

def main():
    model_files = find_best_models(SAVED_MODELS_DIR)
    if not model_files:
        print("No *_best_model.keras files found.")
        sys.exit(1)

    good_dir = TEST_IMAGES_DIR / "good"
    bad_dir = TEST_IMAGES_DIR / "bad"
    good_imgs = sorted([p for p in good_dir.glob("*.*") if p.suffix.lower() in [".jpg", ".jpeg", ".png"]])[:NUM_GOOD]
    bad_imgs = sorted([p for p in bad_dir.glob("*.*") if p.suffix.lower() in [".jpg", ".jpeg", ".png"]])[:NUM_BAD]

    if len(good_imgs) < NUM_GOOD or len(bad_imgs) < NUM_BAD:
        print(f"⚠️ Need at least {NUM_GOOD} good and {NUM_BAD} bad images.")
        sys.exit(1)

    models_info = []
    for model_path in model_files:
        model_name = model_path.parent.name
        print(f"\nLoading model: {model_name}")
        model = load_model(str(model_path))
        preds = []
        for p in list(good_imgs) + list(bad_imgs):
            img_tensor = load_and_prepare_image(p, IMG_SIZE)
            idx, conf, raw = predict_with_model(model, img_tensor)
            pred_label = CLASS_MAPPING.get(idx, str(idx))
            true_label = p.parent.name.lower()
            correct = (pred_label.lower() == true_label)
            preds.append({
                "img_path": str(p),
                "true_label": true_label,
                "pred_label": pred_label,
                "confidence": conf,
                "correct": correct
            })
        models_info.append((model_name, preds))

    out_png = RESULTS_DIR / "model_comparison_4.png"
    build_visual_grid(models_info, out_png)

    out_docx = RESULTS_DIR / "report_4.docx"
    create_docx_report(models_info, out_docx, out_png)

if __name__ == "__main__":
    main()
